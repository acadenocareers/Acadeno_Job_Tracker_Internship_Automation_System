import time
import re
import requests
from bs4 import BeautifulSoup
from docx import Document
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import os


chrome_options = Options()
chrome_options.add_argument("headless")
chrome_options.add_argument("disable-gpu")
chrome_options.add_argument("no-sandbox")
chrome_options.add_argument("disable-dev-shm-usage")
chrome_options.add_argument("user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36")

driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=chrome_options)


TECHNICAL_ROLES = [
    # Data Roles
    "data scientist", "data science", "data analyst", "data analytics", 
    "business intelligence", "bi developer", "data engineer", "etl", 
    "power bi", "tableau", "machine learning", "ml", "ai", "artificial intelligence",
    "analytics",
    
    # Python Roles
    "python", "django", "flask", "fastapi",
    
    # Full Stack & Web Development
    "full stack", "fullstack", "web developer", "mern", "mean",
    
    # Frontend Technologies
    "frontend", "front end", "react", "angular", "vue", "javascript", 
    "typescript", "html", "css", "bootstrap", "tailwind", "ui developer", 
    "ux developer",
    
    # Mobile Development
    "flutter", "react native", "mobile developer",
    
    # Internships
    "intern", "internship", "trainee"
]

# Roles to EXCLUDE - STRICT
EXCLUDE_ROLES = ["php", "laravel", "wordpress", "joomla", "drupal", "magento", ".net", "c#"]


def fetch_infopark_jobs():
    print("Fetching jobs from Infopark...")
    jobs = []
    page = 1
    max_pages = 10
    
    while page <= max_pages:
        try:
            url = f"https://infopark.in/companies/job-search?page={page}"
            print(f"Scraping page {page}: {url}")
            
            driver.get(url)
            time.sleep(2)
            
            # Parse the table
            soup = BeautifulSoup(driver.page_source, 'html.parser')
            table = soup.find('table')
            
            if not table:
                print("No table found, stopping...")
                break
                
            rows = table.find_all('tr')
            if len(rows) <= 1:
                print("No job rows found, stopping...")
                break
            
            page_jobs_count = 0
            for row in rows[1:]:
                try:
                    cols = row.find_all('td')
                    if len(cols) >= 3:
                        date = cols[0].get_text(strip=True)
                        title = cols[1].get_text(strip=True)
                        company = cols[2].get_text(strip=True)
                        
                        job_link = ""
                        link_element = row.find('a', href=True)
                        if link_element:
                            job_link = link_element['href']
                            if not job_link.startswith('http'):
                                job_link = f"https://infopark.in{job_link}" if job_link.startswith('/') else f"https://infopark.in/{job_link}"
                        
                        title_lower = title.lower()
                        
                        if any(excluded in title_lower for excluded in EXCLUDE_ROLES):
                            continue
                            
                        if any(tech_role in title_lower for tech_role in TECHNICAL_ROLES):
                            experience = "0-2 years"
                            
                            exp_patterns = [
                                r'(\d+\s*-\s*\d+\s*(?:years?|yrs?))',
                                r'(\d+\s*(?:years?|yrs?))',
                            ]
                            
                            for pattern in exp_patterns:
                                match = re.search(pattern, title_lower)
                                if match:
                                    exp_text = match.group(1).strip()
                                    experience = exp_text
                                    break
                            
                            if "intern" in title_lower or "trainee" in title_lower:
                                experience = "Intern"
                            
                            jobs.append({
                                "park": "Infopark",
                                "title": title,
                                "company": company,
                                "experience": experience,
                                "date": date,
                                "location": "Infopark, Kochi",
                                "link": job_link
                            })
                            page_jobs_count += 1
                            
                except Exception:
                    continue
            
            print(f"Page {page}: Found {page_jobs_count} technical jobs")
            
            if page_jobs_count == 0:
                print("No technical jobs found on this page, stopping...")
                break
                
            page += 1
            
        except Exception as e:
            print(f"Error on page {page}: {e}")
            break
    
    print(f" Found {len(jobs)} technical jobs from Infopark.")
    return jobs

def fetch_technopark_jobs():
    print("Fetching jobs from Technopark...")
    jobs = []
    try:
        driver.get("https://technopark.org/job-search")
        time.sleep(4)
        
        soup = BeautifulSoup(driver.page_source, 'html.parser')
        
        job_elements = soup.find_all(['div', 'li', 'article', 'tr'])
        
        for elem in job_elements:
            try:
                text = elem.get_text(strip=True)
                if len(text) < 20:
                    continue
                    
                text_lower = text.lower()
                
                if any(excluded in text_lower for excluded in EXCLUDE_ROLES):
                    continue
                    
                if any(tech_role in text_lower for tech_role in TECHNICAL_ROLES):
                    lines = [line.strip() for line in text.split('\n') if line.strip()]
                    title = lines[0] if lines else text[:50]
                    
                    if len(title) < 5:
                        continue
                        
                    company = "Not specified"
                    job_link = ""
                    
                    link_element = elem.find('a', href=True)
                    if link_element:
                        job_link = link_element['href']
                        if not job_link.startswith('http'):
                            job_link = f"https://technopark.org{job_link}" if job_link.startswith('/') else f"https://technopark.org/{job_link}"
                    
                    for line in lines:
                        if any(term in line.lower() for term in ["pvt", "ltd", "inc", "technologies", "solutions"]):
                            company = line
                            break
                    
                    experience = "0-2 years"
                    if "intern" in text_lower:
                        experience = "Intern"
                    
                    jobs.append({
                        "park": "Technopark",
                        "title": title,
                        "company": company,
                        "experience": experience,
                        "date": "N/A",
                        "location": "Technopark, Trivandrum",
                        "link": job_link
                    })
                    
            except Exception:
                continue
                
    except Exception as e:
        print(f" Technopark error: {e}")
    
    unique_jobs = []
    seen = set()
    for job in jobs:
        identifier = job['title'].lower()
        if identifier not in seen:
            seen.add(identifier)
            unique_jobs.append(job)
    
    print(f" Found {len(unique_jobs)} technical jobs from Technopark.")
    return unique_jobs

def fetch_cyberpark_jobs():
    print("Fetching jobs from Cyberpark...")
    jobs = []
    try:
        driver.get("https://cyberparkkerala.org/careers")
        time.sleep(4)
        
        all_text = driver.find_element(By.TAG_NAME, "body").text
        lines = [line.strip() for line in all_text.split('\n') if line.strip()]
        
        seen_titles = set()
        
        for line in lines:
            line_lower = line.lower()
            
            if any(excluded in line_lower for excluded in EXCLUDE_ROLES):
                continue
                
            if any(tech_role in line_lower for tech_role in TECHNICAL_ROLES):
                is_job_title = (
                    len(line) > 10 and 
                    len(line) < 100 and
                    not line.startswith('©') and
                    not line.startswith('http') and
                    any(term in line_lower for term in ["developer", "engineer", "analyst", "specialist", "intern", "trainee"])
                )
                
                if is_job_title and line not in seen_titles:
                    seen_titles.add(line)
                    
                    experience = "0-2 years"
                    if "intern" in line_lower or "trainee" in line_lower:
                        experience = "Intern"
                        
                    jobs.append({
                        "park": "Cyberpark",
                        "title": line,
                        "company": "Cyberpark Company",
                        "experience": experience,
                        "date": "N/A",
                        "location": "Cyberpark, Kozhikode",
                        "link": "https://cyberparkkerala.org/careers"  # General careers page
                    })
                        
    except Exception as e:
        print(f"Cyberpark error: {e}")
    
    print(f" Found {len(jobs)} technical jobs from Cyberpark (PHP/Laravel excluded).")
    return jobs

def filter_jobs_by_experience(jobs):
    filtered_jobs = []
    
    for job in jobs:
        experience = job["experience"].lower()
        title_lower = job["title"].lower()
        
        if "intern" in title_lower or "trainee" in title_lower:
            filtered_jobs.append(job)
            continue
            
        valid_patterns = [
            r'^0\s*-\s*[12]\s*(?:years?|yrs?)',
            r'^[0-2]\s*(?:years?|yrs?)',
            r'fresher',
            r'entry[-\s]level',
            r'^0-2',
            r'^1-2',
            r'^0-1'
        ]
        
        invalid_patterns = [
            r'[3-9]\s*[-+]\s*[0-9]',
            r'[3-9]\s*(?:years?|yrs?)',
            r'[3-9]\+',
            r'senior',
            r'lead',
            r'architect',
            r'manager',
            r'experienced'
        ]
        
        is_valid = any(re.search(pattern, experience, re.IGNORECASE) for pattern in valid_patterns)
        is_invalid = any(re.search(pattern, experience, re.IGNORECASE) for pattern in invalid_patterns)
        
        if experience == "0-2 years" and not is_invalid:
            filtered_jobs.append(job)
        elif is_valid and not is_invalid:
            filtered_jobs.append(job)
    
    print(f" Filtered to {len(filtered_jobs)} jobs with 0-2 years experience.")
    return filtered_jobs

def save_jobs_to_word(jobs, filename="technical_jobs_0_2_years.docx"):
    if not jobs:
        print(" No technical jobs found matching your criteria.")
        return
        
    try:
        if os.path.exists(filename):
            os.rename(filename, f"backup_{filename}")
    except:
        pass
    
    try:
        doc = Document()
        doc.add_heading("Technical Job Openings in Kerala IT Parks", level=1)
        doc.add_paragraph("Fetched from Infopark, Technopark, and Cyberpark")
        
        doc.add_paragraph("   • Data Science & Analytics Roles")
        doc.add_paragraph("   • Python Development") 
        doc.add_paragraph("   • Full Stack Development")
        doc.add_paragraph("   • Frontend Technologies (React, Angular, Vue, CSS)")
        doc.add_paragraph("   • Flutter & Mobile Development")
        doc.add_paragraph("   • Internships")
        doc.add_paragraph("   • Experience: 0-2 years only")
        doc.add_paragraph("   • Excluded: PHP, Laravel, .NET, C#, Senior roles")
        doc.add_paragraph(f"Total matching jobs found: {len(jobs)}\n")
        
        # Group by IT park
        parks = {}
        for job in jobs:
            park = job["park"]
            if park not in parks:
                parks[park] = []
            parks[park].append(job)
        
        for park, park_jobs in parks.items():
            doc.add_heading(park, level=2)
            
            for idx, job in enumerate(park_jobs, 1):
                p = doc.add_paragraph()
                p.add_run(f"{idx}. {job['title']}\n").bold = True
                p.add_run(f"    Company: {job['company']}\n")
                p.add_run(f"    Location: {job['location']}\n")
                p.add_run(f"    Experience: {job['experience']}\n")
                if job['date'] != "N/A":
                    p.add_run(f"    Date: {job['date']}\n")
                
                # Add job link if available
                if job['link']:
                    p.add_run(f"    Job Link: {job['link']}\n")
                else:
                    p.add_run(f"    Job Link: Not available\n")
                    
                p.add_run(f"    Source: {job['park']}\n")
        
        doc.save(filename)
        print(f" Technical job listings saved to {filename}")
    except PermissionError:
        alt_filename = f"jobs_{int(time.time())}.docx"
        doc.save(alt_filename)
        print(f" Permission denied for {filename}, saved as {alt_filename} instead")
    except Exception as e:
        print(f" Error saving file: {e}")

def display_jobs_in_console(jobs):
    if not jobs:
        print("No jobs to display")
        return
        
    parks = {}
    for job in jobs:
        park = job["park"]
        if park not in parks:
            parks[park] = []
        parks[park].append(job)
    
    print(f"\n FINAL TECHNICAL JOBS FOUND ({len(jobs)} total):")
    
    for park, park_jobs in parks.items():
        print(f"\n {park.upper()} ({len(park_jobs)} jobs):")
        print("-" * 60)
        for job in park_jobs:
            print(f"  • {job['title']}")
            print(f"    Company: {job['company']} | Experience: {job['experience']}")
            if job['link']:
                print(f"    Link: {job['link']}")
            if job['date'] != "N/A":
                print(f"    Date: {job['date']}")


def main():
    print("   • Data Science, Data Analytics, ML, AI")
    print("   • Python, Django, Flask")
    print("   • Full Stack, Web Development")
    print("   • Frontend: React, Angular, Vue, JavaScript, CSS")
    print("   • Mobile: Flutter, React Native")
    print("   • Internships & Trainee positions")
    print("   • Experience: 0-2 years ONLY")
  
    
    all_jobs = []

    # Fetch from all parks
    print("\n Scanning Infopark for Python, Data Science, Intern roles...")
    infopark_jobs = fetch_infopark_jobs()
    all_jobs.extend(infopark_jobs)
    
    print("\n Scanning Technopark for technical roles...")
    technopark_jobs = fetch_technopark_jobs()
    all_jobs.extend(technopark_jobs)
    
    cyberpark_jobs = fetch_cyberpark_jobs()
    all_jobs.extend(cyberpark_jobs)

    # Close browser
    driver.quit()

    if all_jobs:
        print(f"\n Total technical jobs found before filtering: {len(all_jobs)}")
        
        # Show what was found before filtering
        display_jobs_in_console(all_jobs)
        
        # Filter by experience
        filtered_jobs = filter_jobs_by_experience(all_jobs)
        
        if filtered_jobs:
            print(f"\n FINAL RESULT: {len(filtered_jobs)} jobs match all criteria!")
            display_jobs_in_console(filtered_jobs)
            
            # Save to Word
            save_jobs_to_word(filtered_jobs)
            
        else:
            print(" No jobs found matching both technical role AND 0-2 years experience criteria.")
            print(" Saving all technical jobs to document for reference...")
            save_jobs_to_word(all_jobs, "all_technical_jobs.docx")
            
    else:
        print(" No technical jobs found at all.")

if __name__ == "__main__":
    main()